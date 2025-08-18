import os
import logging
from typing import Optional
import PyPDF2
from docx import Document
from PIL import Image
import io
import base64

logger = logging.getLogger(__name__)

class FileParser:
    """
    Parser for extracting text from various file formats
    """
    
    def __init__(self):
        self.supported_extensions = {'.pdf', '.docx', '.txt', '.png', '.jpg', '.jpeg', '.webp'}
        # Gemini Vision optimization settings
        self.max_image_width = 1024  # Max width for Gemini Vision
        self.max_image_height = 1024  # Max height for Gemini Vision
        self.quality = 85  # JPEG quality (1-100)
        self.max_file_size_mb = 4  # Max file size in MB for Gemini
    
    async def extract_text(self, file_path: str) -> str:
        """
        Extract text from a file based on its extension
        """
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension not in self.supported_extensions:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            if file_extension == '.pdf':
                return await self._extract_from_pdf(file_path)
            elif file_extension == '.docx':
                return await self._extract_from_docx(file_path)
            elif file_extension == '.txt':
                return await self._extract_from_txt(file_path)
            elif file_extension in {'.png', '.jpg', '.jpeg', '.webp'}:
                return await self._extract_from_image(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    async def _extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file
        """
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)
            
            extracted_text = '\n\n'.join(text_content)
            logger.info(f"Extracted text from PDF: {len(extracted_text)} characters")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error extracting from PDF {file_path}: {str(e)}")
            raise
    
    async def _extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file
        """
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            extracted_text = '\n\n'.join(text_content)
            logger.info(f"Extracted text from DOCX: {len(extracted_text)} characters")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error extracting from DOCX {file_path}: {str(e)}")
            raise
    
    async def _extract_from_txt(self, file_path: str) -> str:
        """
        Extract text from TXT file
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            logger.info(f"Extracted text from TXT: {len(text_content)} characters")
            return text_content
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text_content = file.read()
                logger.info(f"Extracted text from TXT (latin-1): {len(text_content)} characters")
                return text_content
            except Exception as e:
                logger.error(f"Error extracting from TXT {file_path}: {str(e)}")
                raise
        except Exception as e:
            logger.error(f"Error extracting from TXT {file_path}: {str(e)}")
            raise

    async def _extract_from_image(self, file_path: str) -> str:
        """
        Extract text from image using Gemini Vision API with optimized preprocessing
        """
        try:
            # Preprocess image for optimal Gemini Vision usage
            optimized_image_data = await self._preprocess_image_for_gemini(file_path)
            
            # Import Gemini client here to avoid circular imports
            from services.gemini_client import GeminiClient
            
            gemini_client = GeminiClient()
            
            # Create a prompt for image analysis
            prompt = """Please analyze this image and extract any text content you can find. 
            If there's no text, describe what you see in the image in detail.
            If it's a document, extract all the text content.
            If it's a photo, describe the scene, objects, people, or animals visible.
            Please provide a comprehensive description."""
            
            # Use Gemini Vision to extract text/description
            response = await gemini_client.generate_response_with_image(
                prompt=prompt,
                image_data=optimized_image_data
            )
            
            logger.info(f"Extracted content from image using Gemini Vision: {len(response)} characters")
            return response
            
        except Exception as e:
            logger.error(f"Error extracting from image {file_path}: {str(e)}")
            raise
    
    async def _preprocess_image_for_gemini(self, file_path: str) -> str:
        """
        Preprocess image to optimize for Gemini Vision API:
        - Resize to reasonable dimensions
        - Compress to reduce file size
        - Convert to base64
        """
        try:
            # Open image
            with Image.open(file_path) as img:
                # Convert to RGB if necessary (Gemini prefers RGB)
                if img.mode in ('RGBA', 'LA', 'P'):
                    img = img.convert('RGB')
                
                # Get original dimensions
                original_width, original_height = img.size
                logger.info(f"Original image size: {original_width}x{original_height}")
                
                # Calculate new dimensions while maintaining aspect ratio
                if original_width > self.max_image_width or original_height > self.max_image_height:
                    # Calculate scaling factor
                    scale = min(self.max_image_width / original_width, 
                               self.max_image_height / original_height)
                    
                    new_width = int(original_width * scale)
                    new_height = int(original_height * scale)
                    
                    # Resize image
                    img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                    logger.info(f"Resized image to: {new_width}x{new_height}")
                
                # Save to bytes with compression
                img_bytes = io.BytesIO()
                
                # Determine format and save with optimization
                if img.format == 'JPEG' or file_path.lower().endswith(('.jpg', '.jpeg')):
                    img.save(img_bytes, format='JPEG', quality=self.quality, optimize=True)
                else:
                    # Convert to JPEG for better compression
                    img.save(img_bytes, format='JPEG', quality=self.quality, optimize=True)
                
                img_bytes.seek(0)
                
                # Convert to base64
                base64_data = base64.b64encode(img_bytes.getvalue()).decode('utf-8')
                
                # Check file size
                file_size_mb = len(img_bytes.getvalue()) / (1024 * 1024)
                logger.info(f"Processed image size: {file_size_mb:.2f} MB")
                
                if file_size_mb > self.max_file_size_mb:
                    logger.warning(f"Image still large ({file_size_mb:.2f} MB), consider further compression")
                
                return base64_data
                
        except Exception as e:
            logger.error(f"Error preprocessing image {file_path}: {str(e)}")
            raise
    
    def get_supported_extensions(self) -> set:
        """
        Get list of supported file extensions
        """
        return self.supported_extensions.copy() 